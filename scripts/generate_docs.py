from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "word"
ASSETS_DIR = ROOT / "docs" / "assets"


def _font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    # Best-effort: use a common system font, fall back to default bitmap font.
    for name in [
        "C:\\Windows\\Fonts\\segoeui.ttf",
        "C:\\Windows\\Fonts\\arial.ttf",
    ]:
        p = Path(name)
        if p.exists():
            return ImageFont.truetype(str(p), size=size)
    return ImageFont.load_default()


def _save_png(path: Path, img: Image.Image) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")
    return path


def make_tech_stack_diagram(path: Path) -> Path:
    w, h = 1400, 850
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = _font(44)
    box_f = _font(28)

    d.text((40, 30), "Technology stack (offline)", fill="black", font=title_f)

    def box(x, y, bw, bh, text, fill="#f1f5f9"):
        d.rounded_rectangle([x, y, x + bw, y + bh], radius=18, outline="#0f172a", width=3, fill=fill)
        tw, th = d.multiline_textbbox((0, 0), text, font=box_f, spacing=6)[2:]
        d.multiline_text(
            (x + (bw - tw) / 2, y + (bh - th) / 2),
            text,
            fill="#0f172a",
            font=box_f,
            align="center",
            spacing=6,
        )

    box(60, 130, 520, 160, "Desktop UI\nTauri 2 + React + TS", fill="#dbeafe")
    box(820, 130, 520, 160, "Worker\nPython + faster-whisper", fill="#dcfce7")
    box(60, 390, 520, 160, "Database\nSQLite (job history)", fill="#fef9c3")
    box(820, 390, 520, 160, "Media decode\nFFmpeg → 16 kHz WAV", fill="#fee2e2")
    box(440, 640, 520, 160, "Model runtime\nCTranslate2 (CPU/GPU)", fill="#e9d5ff")

    def arrow(x1, y1, x2, y2):
        d.line([x1, y1, x2, y2], fill="#0f172a", width=5)
        ang = math.atan2(y2 - y1, x2 - x1)
        ah = 18
        p1 = (x2 - ah * math.cos(ang - 0.35), y2 - ah * math.sin(ang - 0.35))
        p2 = (x2 - ah * math.cos(ang + 0.35), y2 - ah * math.sin(ang + 0.35))
        d.polygon([ (x2, y2), p1, p2 ], fill="#0f172a")

    arrow(580, 210, 820, 210)
    arrow(320, 290, 320, 390)
    arrow(1080, 290, 1080, 390)
    arrow(1080, 550, 700, 640)

    return _save_png(path, img)


def make_dataflow_diagram(path: Path) -> Path:
    w, h = 1600, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = _font(44)
    box_f = _font(28)
    small_f = _font(22)
    d.text((40, 30), "Dataflow diagram", fill="black", font=title_f)

    def box(x, y, bw, bh, title, body, fill="#f8fafc"):
        d.rounded_rectangle([x, y, x + bw, y + bh], radius=16, outline="#0f172a", width=3, fill=fill)
        d.text((x + 18, y + 14), title, fill="#0f172a", font=box_f)
        d.multiline_text((x + 18, y + 64), body, fill="#334155", font=small_f, spacing=6)

    box(80, 160, 380, 180, "UI", "Select files\nStart queue\nMenu: Export/Copy", fill="#dbeafe")
    box(520, 160, 420, 180, "FFmpeg", "Decode media\nExtract mono 16 kHz WAV", fill="#fee2e2")
    box(1000, 160, 520, 180, "ASR (faster-whisper)", "VAD + decode\nSegments with timestamps", fill="#dcfce7")
    box(520, 420, 420, 180, "SQLite", "Store jobs\nProgress, errors\nSegments JSON", fill="#fef9c3")
    box(1000, 420, 520, 180, "Export", "SRT / VTT / TXT\nSaved to chosen path", fill="#e9d5ff")

    def arrow(x1, y1, x2, y2, label=None):
        d.line([x1, y1, x2, y2], fill="#0f172a", width=5)
        ang = math.atan2(y2 - y1, x2 - x1)
        ah = 18
        p1 = (x2 - ah * math.cos(ang - 0.35), y2 - ah * math.sin(ang - 0.35))
        p2 = (x2 - ah * math.cos(ang + 0.35), y2 - ah * math.sin(ang + 0.35))
        d.polygon([(x2, y2), p1, p2], fill="#0f172a")
        if label:
            d.text(((x1 + x2) / 2 - 40, (y1 + y2) / 2 - 30), label, fill="#0f172a", font=small_f)

    arrow(460, 250, 520, 250, "paths")
    arrow(940, 250, 1000, 250, "wav")
    arrow(1210, 340, 1210, 420, "segments")
    arrow(940, 510, 1000, 510, "read")
    arrow(730, 340, 730, 420, "status")
    arrow(520, 510, 460, 510, "jobs")

    return _save_png(path, img)


def make_ui_mock_screenshot(path: Path) -> Path:
    w, h = 1600, 900
    img = Image.new("RGB", (w, h), "#f8fafc")
    d = ImageDraw.Draw(img)
    title_f = _font(34)
    body_f = _font(22)

    d.rounded_rectangle([40, 40, w - 40, h - 40], radius=20, outline="#cbd5e1", width=4, fill="white")
    d.text((70, 70), "Video Transcript Extractor (illustrative)", fill="#0f172a", font=title_f)
    d.text((70, 115), "Offline transcription • Local processing", fill="#334155", font=body_f)

    # Header chip
    d.rounded_rectangle([1180, 70, w - 70, 118], radius=22, outline="#e2e8f0", width=2, fill="#f1f5f9")
    d.text((1200, 83), "Developed By: Ninad K.", fill="#0f172a", font=body_f)

    # Toolbar
    d.rounded_rectangle([70, 160, w - 70, 220], radius=14, outline="#e2e8f0", width=2, fill="#f1f5f9")
    d.text((90, 176), "Add videos   Start queue   Cancel", fill="#0f172a", font=body_f)
    d.rounded_rectangle([1050, 168, 1340, 212], radius=999, outline="#e2e8f0", width=2, fill="#ffffff")
    d.text((1070, 178), "Total: 42%  ▓▓▓▓▓░░░░░", fill="#0f172a", font=body_f)
    d.text((1400, 176), "Help", fill="#0f172a", font=body_f)

    # Menu hint row
    d.text((70, 230), "Menu: File | Edit | View | Export | Help", fill="#64748b", font=body_f)

    # Left queue panel
    d.rounded_rectangle([70, 270, 1020, h - 70], radius=16, outline="#e2e8f0", width=2, fill="white")
    d.text((95, 295), "Queue", fill="#0f172a", font=title_f)
    for i in range(5):
        y = 355 + i * 105
        d.rounded_rectangle([95, y, 995, y + 85], radius=14, outline="#e2e8f0", width=2, fill="#ffffff")
        d.text((120, y + 14), f"D:\\\\videos\\\\sample_{i+1}.mp4", fill="#0f172a", font=body_f)
        d.text((120, y + 48), "status: queued    progress: 0%   ░░░░░░░░░░", fill="#475569", font=body_f)

    # Settings panel
    d.rounded_rectangle([1050, 270, w - 70, 720], radius=16, outline="#e2e8f0", width=2, fill="white")
    d.text((1075, 295), "Settings", fill="#0f172a", font=title_f)
    d.text((1075, 360), "FFmpeg path (optional)", fill="#0f172a", font=body_f)
    d.rounded_rectangle([1075, 400, w - 95, 440], radius=10, outline="#cbd5e1", width=2, fill="#ffffff")
    d.text((1088, 410), "ffmpeg on PATH if empty", fill="#94a3b8", font=body_f)
    d.text((1075, 470), "Model cache directory", fill="#0f172a", font=body_f)
    d.rounded_rectangle([1075, 510, w - 95, 550], radius=10, outline="#cbd5e1", width=2, fill="#ffffff")
    d.text((1088, 520), "System default if empty", fill="#94a3b8", font=body_f)
    d.text((1075, 580), "Default quality", fill="#0f172a", font=body_f)
    d.rounded_rectangle([1075, 620, 1290, 660], radius=10, outline="#cbd5e1", width=2, fill="#ffffff")
    d.text((1090, 630), "Final", fill="#0f172a", font=body_f)
    d.text((1320, 580), "Default language", fill="#0f172a", font=body_f)
    d.rounded_rectangle([1320, 620, w - 95, 660], radius=10, outline="#cbd5e1", width=2, fill="#ffffff")
    d.text((1335, 630), "Auto", fill="#0f172a", font=body_f)

    return _save_png(path, img)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)


def doc_user_guide(out_path: Path, screenshot_path: Path) -> None:
    doc = Document()
    doc.add_heading("Video Transcript Extractor — User Guide", level=0)
    p = doc.add_paragraph("This document describes how to use the app to extract transcripts offline.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("1. Install prerequisites", level=1)
    doc.add_paragraph("Install Python 3.10+, FFmpeg, and (for Windows) Visual Studio Build Tools + C++ workload.")
    doc.add_paragraph("Install Node.js + npm to run the desktop UI during development.")

    doc.add_heading("2. Add videos and start the queue", level=1)
    doc.add_paragraph("Use File → Add videos… or click “Add videos”, choose one or more files, then click “Start queue”.")
    doc.add_paragraph("Quality and language defaults are configured under Settings.")

    doc.add_paragraph("Illustrative UI screenshot:").runs[0].bold = True
    doc.add_picture(str(screenshot_path), width=Inches(6.7))

    doc.add_heading("3. Review and export", level=1)
    doc.add_paragraph("After completion, review segments and export SRT, VTT, or TXT.")

    doc.add_heading("Troubleshooting", level=1)
    doc.add_paragraph("If FFmpeg is not found, set its path in Settings.")
    doc.add_paragraph("If you see CUDA out-of-memory, switch to Draft or free VRAM.")

    doc.add_heading("Help menu", level=1)
    doc.add_paragraph("Use Help → User guide from the app menu to open the in-app Help screen.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def doc_architecture(out_path: Path, stack_png: Path, flow_png: Path) -> None:
    doc = Document()
    doc.add_heading("Video Transcript Extractor — Architecture", level=0)
    doc.add_paragraph("Offline transcription architecture and diagrams.")

    doc.add_heading("Technology stack", level=1)
    doc.add_paragraph("Desktop UI: Tauri 2 + React + TypeScript.")
    doc.add_paragraph("Worker: Python + faster-whisper (CTranslate2).")
    doc.add_paragraph("Media decode: FFmpeg.")
    doc.add_paragraph("Persistence: SQLite (job history, segments JSON).")

    doc.add_picture(str(stack_png), width=Inches(6.8))

    doc.add_heading("Dataflow", level=1)
    doc.add_paragraph("Files are selected in the UI and processed by the worker; progress and results are stored and displayed.")
    doc.add_picture(str(flow_png), width=Inches(6.8))

    doc.add_heading("Notes", level=1)
    doc.add_paragraph("Draft uses medium; Final uses large-v3.")
    doc.add_paragraph("All processing is performed on-device; no audio is uploaded.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    stack_png = make_tech_stack_diagram(ASSETS_DIR / "technology_stack.png")
    flow_png = make_dataflow_diagram(ASSETS_DIR / "dataflow.png")
    ui_png = make_ui_mock_screenshot(ASSETS_DIR / "ui_screenshot_illustrative.png")

    doc_user_guide(OUT_DIR / "User_Guide.docx", ui_png)
    doc_architecture(OUT_DIR / "Architecture_and_Diagrams.docx", stack_png, flow_png)

    print(f"Wrote: {OUT_DIR / 'User_Guide.docx'}")
    print(f"Wrote: {OUT_DIR / 'Architecture_and_Diagrams.docx'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

